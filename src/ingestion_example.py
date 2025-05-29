import os
import asyncio
from typing import List
from docx import Document as DocxDocument
from llama_index.core import Document, Settings
from llama_index.core.node_parser import SimpleNodeParser
from llama_index.core.ingestion import IngestionPipeline
from llama_index.embeddings.huggingface import HuggingFaceEmbedding
from easyrag.utils import get_yaml_data

def read_docx(file_path: str) -> str:
    """读取 docx 文件内容"""
    doc = DocxDocument(file_path)
    full_text = []
    
    # 读取正文段落
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():  # 跳过空段落
            full_text.append(paragraph.text)
    
    # 读取表格内容
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                if cell.text.strip():  # 跳过空单元格
                    row_text.append(cell.text.strip())
            if row_text:  # 只添加非空行
                full_text.append(" | ".join(row_text))
    
    return "\n".join(full_text)

async def load_documents(data_path: str) -> List[Document]:
    """加载文档"""
    documents = []
    for root, _, files in os.walk(data_path):
        for file in files:
            file_path = os.path.join(root, file)
            try:
                if file.endswith('.docx'):
                    # 处理 Word 文档
                    text = read_docx(file_path)
                elif file.endswith(('.txt', '.md')):
                    # 处理文本文件
                    with open(file_path, 'r', encoding='utf-8') as f:
                        text = f.read()
                else:
                    # 跳过不支持的文件类型
                    continue
                
                # 创建文档对象
                if text.strip():  # 只处理非空文档
                    doc = Document(
                        text=text,
                        metadata={
                            'file_name': file,
                            'file_path': file_path,
                            'file_type': file.split('.')[-1],
                        }
                    )
                    documents.append(doc)
                    print(f"Successfully loaded: {file}")
            except Exception as e:
                print(f"Error loading {file}: {str(e)}")
                continue
    
    return documents

async def main():
    # 加载配置
    config = get_yaml_data("configs/easyrag.yaml")
    
    # 初始化 Embedding 模型
    embedding = HuggingFaceEmbedding(
        model_name="BAAI/bge-large-zh-v1.5",
        cache_folder="./cache/models",
    )
    
    # 设置全局 embedding 模型
    Settings.embed_model = embedding
    
    # 创建文档解析器
    node_parser = SimpleNodeParser.from_defaults(
        chunk_size=512,
        chunk_overlap=50,
    )
    
    # 创建处理管道（只使用文档分块）
    pipeline = IngestionPipeline(
        transformations=[
            node_parser,  # 文档分块
        ]
    )
    
    try:
        # 加载文档
        data_path = "/mnt/sda/users/andy/project/Rag/BUAA/EasyRAG/data/赛题制度文档"
        print("\nLoading documents...")
        documents = await load_documents(data_path)
        print(f"\nLoaded {len(documents)} documents")
        
        # 处理前3个文档
        documents = documents[:3]
        print(f"Processing first {len(documents)} documents as a test...")
        
        # 运行处理管道
        print("\nProcessing documents...")
        nodes = await pipeline.arun(
            documents=documents,
            show_progress=True,  # 显示进度条
        )
        
        print(f"\nProcessed {len(nodes)} nodes")
        
        # 打印处理结果示例
        for i, node in enumerate(nodes[:2]):  # 只打印前两个节点的信息
            print(f"\nNode {i + 1}:")
            print(f"Text: {node.text[:200]}...")  # 显示前200个字符
            print("\nMetadata:")
            for key, value in node.metadata.items():
                print(f"  {key}: {value}")
    
    except Exception as e:
        print(f"Error: {str(e)}")
        import traceback
        print(f"Traceback: {traceback.format_exc()}")

if __name__ == "__main__":
    asyncio.run(main()) 