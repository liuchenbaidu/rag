import os
import json
from docx import Document
from tqdm import tqdm
import re


def extract_document_properties(doc):
    """提取文档的属性信息"""
    properties = {}
    core_properties = doc.core_properties
    if core_properties:
        if core_properties.title:
            properties['title'] = core_properties.title
        if core_properties.subject:
            properties['subject'] = core_properties.subject
        if core_properties.author:
            properties['author'] = core_properties.author
        if core_properties.keywords:
            properties['keywords'] = core_properties.keywords
    return properties


def clean_text(text):
    """清理文本，去除多余的空白字符但保留基本格式"""
    if not text:
        return ""
    # 替换多个空白字符为单个空格
    text = re.sub(r'\s+', ' ', text)
    # 清理空白字符但保留段落之间的分隔
    text = text.strip()
    return text


def process_docx(docx_path):
    """处理单个docx文件，提取文本内容"""
    try:
        doc = Document(docx_path)
        content = []
        
        # 提取文档属性
        properties = extract_document_properties(doc)
        if properties:
            content.append("### 文档信息")
            for key, value in properties.items():
                if value:
                    content.append(f"{key}: {value}")
            content.append("")  # 添加空行分隔
        
        # 提取所有段落的文本
        for paragraph in doc.paragraphs:
            text = clean_text(paragraph.text)
            if text and not text.isspace():
                # 检查段落的样式
                if hasattr(paragraph, 'style') and paragraph.style:
                    style_name = paragraph.style.name
                    # 如果是标题样式，添加特殊标记
                    if 'Heading' in style_name or '标题' in style_name:
                        text = f"### {text}"
                    elif style_name != 'Normal':  # 如果不是普通样式，添加样式信息
                        content.append(f"[样式: {style_name}]")
                content.append(text)
        
        # 提取表格中的文本
        for table in doc.tables:
            table_content = []
            header_row = []
            
            # 首先处理表头
            if table.rows:
                for cell in table.rows[0].cells:
                    header_text = clean_text(cell.text)
                    if header_text:
                        header_row.append(header_text)
            
            if header_row:
                table_content.append("### 表格")
                table_content.append(" | ".join(header_row))
                table_content.append("-" * (len(" | ".join(header_row)) + 2))  # 添加分隔线
            
            # 处理数据行
            for row in table.rows[1:]:
                row_texts = []
                for cell in row.cells:
                    cell_text = clean_text(cell.text)
                    if cell_text:
                        row_texts.append(cell_text)
                    else:
                        row_texts.append(" ")  # 空单元格用空格占位
                if any(text.strip() for text in row_texts):  # 如果行中有非空内容
                    table_content.append(" | ".join(row_texts))
            
            if table_content:
                content.extend(table_content)
                content.append("")  # 在表格后添加空行
        
        # 过滤空行并合并内容
        content = [line for line in content if line.strip()]
        
        if not content:
            print(f"警告: 文档内容为空 {docx_path}")
            # 尝试使用文件名作为标题
            filename = os.path.splitext(os.path.basename(docx_path))[0]
            content = ["### " + filename]
        
        return '\n'.join(content)
    except Exception as e:
        print(f"处理文档失败 {docx_path}: {str(e)}")
        return None


def format_content(content, path):
    """格式化文档内容"""
    if content is None:
        print('空文档: ', path)
        return '<文档为空>\n'

    new_content = []
    last_line = None
    
    for line in content.split('\n'):
        line = line.strip()
        # 跳过完全相同的连续行，但保留标题、表格和样式信息
        if line == last_line and not (line.startswith('###') or '|' in line or line.startswith('[样式:')):
            continue
        last_line = line
        if line:
            new_content.append(line)
    
    new_str = ''
    if args.with_path:
        new_str += '###\n'
        new_str += '文档路径: ' + '/'.join(path) + '\n\n'

    if new_content:
        new_str += '\n'.join(new_content) + '\n'
    else:
        print('空文档: ', path)
        new_str += '<文档为空>\n'
    return new_str


def find_docx_files(directory):
    """递归查找目录下的所有docx文件"""
    docx_files = []
    for root, _, files in os.walk(directory):
        for file in files:
            if file.endswith('.docx') and not file.startswith('~$'):  # 排除临时文件
                docx_files.append(os.path.join(root, file))
    return sorted(docx_files)  # 排序以保持处理顺序一致


def process_directory(input_dir, output_dir):
    """处理指定目录下的所有docx文件"""
    filepath_2_knowpath = {}
    empty_docs = []
    processed_docs = 0
    failed_docs = []
    
    # 获取所有docx文件
    docx_files = find_docx_files(input_dir)
    
    if not docx_files:
        print(f"警告: 在 {input_dir} 中没有找到.docx文件")
        return filepath_2_knowpath

    # 处理每个docx文件
    for input_path in tqdm(docx_files, desc=f"处理文档"):
        try:
            relative_path = os.path.relpath(input_path, input_dir)
            output_path = os.path.join(output_dir, os.path.splitext(relative_path)[0] + '.txt')
            
            # 创建输出目录
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            
            # 处理文档内容
            content = process_docx(input_path)
            
            # 构建文档路径信息
            doc_path = tuple(os.path.splitext(relative_path)[0].split(os.sep))
            
            # 保存处理后的内容
            if content and content.strip() and content.strip() != '<文档为空>':
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(format_content(content, doc_path))
                # 记录文件路径映射
                filepath_2_knowpath[os.path.splitext(relative_path)[0] + '.txt'] = doc_path
                processed_docs += 1
            else:
                empty_docs.append(relative_path)
        except Exception as e:
            print(f"\n处理失败 {relative_path}: {str(e)}")
            failed_docs.append((relative_path, str(e)))
    
    # 打印处理统计
    print(f"\n文档处理统计:")
    print(f"- 总文档数: {len(docx_files)}")
    print(f"- 成功处理: {processed_docs}")
    print(f"- 空文档数: {len(empty_docs)}")
    print(f"- 处理失败: {len(failed_docs)}")
    
    if empty_docs:
        print("\n空文档列表:")
        for doc in empty_docs[:10]:
            print(f"- {doc}")
        if len(empty_docs) > 10:
            print(f"... 还有 {len(empty_docs) - 10} 个空文档")
    
    if failed_docs:
        print("\n处理失败的文档:")
        for doc, error in failed_docs[:10]:
            print(f"- {doc}: {error}")
        if len(failed_docs) > 10:
            print(f"... 还有 {len(failed_docs) - 10} 个失败的文档")
    
    return filepath_2_knowpath


if __name__ == '__main__':
    import argparse
    
    # 获取当前脚本所在目录的上级目录
    base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    
    # 设置默认的输入输出目录
    default_input_dir = os.path.join(base_dir, 'data', '赛题制度文档')
    default_output_dir = os.path.join(base_dir, 'data', 'format_data')
    
    parse = argparse.ArgumentParser(description='将Word文档转换为文本文件的工具')
    parse.add_argument('--with_path', action='store_true', default=False,
                      help='是否在输出文件中包含文档路径信息')
    parse.add_argument('--input_dir', type=str, default=default_input_dir,
                      help=f'输入目录，包含要处理的docx文件 (默认: {default_input_dir})')
    parse.add_argument('--output_dir', type=str, default=default_output_dir,
                      help=f'输出目录，用于保存处理后的文本文件 (默认: {default_output_dir})')
    args = parse.parse_args()
    
    print('处理配置:')
    print(f'- 输入目录: {args.input_dir}')
    print(f'- 输出目录: {args.output_dir}')
    print(f'- 包含路径: {args.with_path}')
    
    # 确保输入目录存在
    if not os.path.exists(args.input_dir):
        print(f"错误: 输入目录不存在: {args.input_dir}")
        exit(1)
    
    # 确保输出目录存在
    os.makedirs(args.output_dir, exist_ok=True)
    
    # 处理文档
    filepath_2_knowpath = process_directory(args.input_dir, args.output_dir)
    
    if not filepath_2_knowpath:
        print("\n错误: 没有成功处理任何文档！")
        exit(1)
    
    # 保存路径映射
    pathmap_file = os.path.join(args.output_dir, "pathmap.json")
    with open(pathmap_file, 'w', encoding='utf-8') as f:
        json.dump(filepath_2_knowpath, f, ensure_ascii=False, indent=4)
    
    print(f"\n处理完成! 路径映射已保存到: {pathmap_file}")
    print(f"共处理了 {len(filepath_2_knowpath)} 个文档") 